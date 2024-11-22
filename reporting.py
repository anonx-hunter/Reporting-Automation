###Reporting Automation :

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

# File paths (update these paths if needed)
excel_path = 'path_to_your_excel_file.xlsx'  # Replace with your Excel file path
word_template_path = 'path_to_your_word_template.docx'  # Replace with your Word template path
output_docx_path = 'output_report.docx'  # Final report path

def generate_vapt_report(excel_file, word_template, output_file):
    # Load Excel file
    workbook = load_workbook(excel_file)
    sheet = workbook.active

    # Load Word template
    template = Document(word_template)

    # Extract data from Excel
    data = []
    header = [cell.value for cell in sheet[1]]  # Assuming the first row contains headers
    for row in sheet.iter_rows(min_row=2, values_only=True):  # Start from the second row
        row_data = dict(zip(header, row))
        data.append(row_data)

    # Create a new Word document for the output
    output_doc = Document()

    # Process each row from the Excel data
    for entry in data:
        for paragraph in template.paragraphs:
            # Copy paragraph content with formatting
            new_paragraph = output_doc.add_paragraph()
            new_paragraph.style = paragraph.style
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                # Apply Arial font and size 11
                new_run.font.name = "Arial"
                new_run.font.size = Pt(11)
                # Preserve bold, italic, underline formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
            
            # Replace placeholders with actual data
            for placeholder, value in entry.items():
                if value is None:
                    value = ""
                if f"{{{{{placeholder}}}}}" in new_paragraph.text:
                    new_paragraph.text = new_paragraph.text.replace(f"{{{{{placeholder}}}}}", str(value))

        # Add a page break between entries (not after the last one)
        if entry != data[-1]:
            output_doc.add_page_break()

    # Save the final report
    output_doc.save(output_file)
    print(f"Report generated successfully at {output_file}")

# Call the function
generate_vapt_report(excel_path, word_template_path, output_docx_path)
